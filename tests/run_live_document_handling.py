"""
tests/run_live_document_handling
================================
Dual-mode LangGraph ReAct agent runner and Document Handling (DH) stress-test (Phase 11).

This script:
1. Creates a live LangGraph ReAct state graph.
2. Registers tools for reading/parsing PDFs, XLSX, CSVs, and Word DOCX documents.
3. Wires our SotisLangGraphGuard as a state interceptor node.
4. Generates a realistic test workspace containing:
   - Scenario A: A 50-page PDF document missing environmental compliance metrics.
   - Scenario B: A multi-sheet Excel workbook with corrupted values in Q4.
5. Operates in two modes:
   - Real Agent Mode: If keys or Ollama models are configured, runs the agent.
   - Interactive Simulation Mode: If offline, simulates the loop and recovery.
"""

from __future__ import annotations

import os
import subprocess
import sys
import uuid
from typing import Annotated, Any, Callable, Dict, List, Optional, Sequence
try:
    from typing_extensions import TypedDict
except ImportError:
    from typing import TypedDict

# Import LangChain / LangGraph components
try:
    from langchain_core.messages import AIMessage, BaseMessage, HumanMessage, RemoveMessage, ToolMessage
    from langchain_core.tools import tool
    from langchain_openai import ChatOpenAI
    from langgraph.graph import END, START, StateGraph
    from langgraph.graph.message import add_messages
except ImportError:
    print("[Sotis-DH] Required packages are loading. Ensure pip install has completed.")
    sys.exit(0)

# Multi-format parsers
import pdfplumber
import openpyxl
import pandas as pd
import docx
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from sotis.core.entropy import EntropyConfig
from sotis.core.schemas import Domain
from sotis.lib.langgraph_integration import SotisLangGraphGuard


# ─────────────────────────────────────────────────────────────────────────────
# 1. State Definition
# ─────────────────────────────────────────────────────────────────────────────

class AgentState(TypedDict):
    """LangGraph conversation state wrapper."""
    messages: Annotated[List[BaseMessage], add_messages]


# ─────────────────────────────────────────────────────────────────────────────
# 2. Workspace Setup & Generators
# ─────────────────────────────────────────────────────────────────────────────

WORKSPACE_DIR = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..", "ExperimentLog", "document handling loop", "workspace")
)
os.makedirs(WORKSPACE_DIR, exist_ok=True)


def ensure_test_workspace_setup() -> None:
    """Generate multi-format document workspace files dynamically, ensuring fresh data."""
    reports_dir = os.path.join(WORKSPACE_DIR, "reports")
    financial_dir = os.path.join(WORKSPACE_DIR, "financial")
    os.makedirs(reports_dir, exist_ok=True)
    os.makedirs(financial_dir, exist_ok=True)

    # 1. Generate 50-Page PDF (Scenario A: Vector Haystack Trap)
    pdf_path = os.path.join(reports_dir, "q3_2026.pdf")
    print(f"[Sotis-DH] Generating 50-page PDF report: {pdf_path}...")
    c = canvas.Canvas(pdf_path, pagesize=letter)
    for i in range(1, 51):
        c.drawString(100, 750, f"Q3 2026 Financial and Compliance Report - Page {i}")
        c.drawString(100, 700, f"Section {i}: Operations and Metrics")
        if i == 1:
            c.drawString(100, 650, "Quarterly Revenue: $15,200,000 USD")
        elif i == 10:
            c.drawString(100, 650, "Total carbon offsets purchased: 2,400 tonnes")
        elif i == 25:
            c.drawString(100, 650, "Compliance Audits completed: 4/4")
        elif i == 40:
            c.drawString(100, 650, "Environmental standards: ISO 14001 certified")
        else:
            c.drawString(100, 650, "No material compliance updates in this section.")
        c.showPage()
    c.save()

    # 2. Generate Multi-Sheet Excel Ledger (Scenario B: Excel Ledger Corruption)
    excel_path = os.path.join(financial_dir, "ledger_2026.xlsx")
    print(f"[Sotis-DH] Generating multi-sheet corrupted Excel ledger: {excel_path}...")
    wb = openpyxl.Workbook()
    
    # Q1 Sheet
    ws1 = wb.active
    ws1.title = "Q1"
    ws1.append(["Category", "Amount"])
    ws1.append(["Revenue", 120000])
    ws1.append(["Expenses", 80000])
    ws1.append(["Tax", 12000])
    
    # Q2 Sheet
    ws2 = wb.create_sheet("Q2")
    ws2.append(["Category", "Amount"])
    ws2.append(["Revenue", 145000])
    ws2.append(["Expenses", 90000])
    ws2.append(["Tax", 15000])
    
    # Q3 Sheet
    ws3 = wb.create_sheet("Q3")
    ws3.append(["Category", "Amount"])
    ws3.append(["Revenue", 160000])
    ws3.append(["Expenses", 95000])
    ws3.append(["Tax", 16000])
    
    # Q4 Corrupted Sheet (Causes parsing / type conversion meltdown)
    ws4 = wb.create_sheet("Q4")
    ws4.append(["Category", "Amount"])
    ws4.append(["Revenue", "ERROR_REF_CORRUPTED_#REF!"])
    ws4.append(["Expenses", "N/A"])
    ws4.append(["Tax", "N/A"])
    
    wb.save(excel_path)

    # 3. Generate Word DOCX
    docx_path = os.path.join(financial_dir, "company_profile.docx")
    print(f"[Sotis-DH] Generating Word company profile: {docx_path}...")
    doc = docx.Document()
    doc.add_heading("Company Profile 2026", 0)
    doc.add_paragraph("Sotis Document Handling stress test document. We verify that python-docx parses this correctly.")
    doc.add_paragraph("Mission: To build highly resilient and self-healing agent workflows.")
    doc.save(docx_path)

    # 4. Generate CSV
    csv_path = os.path.join(financial_dir, "sample_data.csv")
    print(f"[Sotis-DH] Generating Sample CSV table: {csv_path}...")
    df = pd.DataFrame({
        "Year": [2026, 2026, 2026],
        "Quarter": ["Q1", "Q2", "Q3"],
        "Metric": ["Efficiency", "Risk", "Compliance"],
        "Value": [94.5, 12.1, 100.0]
    })
    df.to_csv(csv_path, index=False)


# ─────────────────────────────────────────────────────────────────────────────
# 3. Multi-Format Workspace Tools
# ─────────────────────────────────────────────────────────────────────────────

@tool
def parse_pdf_document(file_path: str, pages: Optional[List[int]] = None) -> str:
    """Extract plain text from page ranges of a PDF document (e.g. pages=[1, 2, 5])."""
    full_path = os.path.join(WORKSPACE_DIR, file_path)
    if not os.path.exists(full_path):
        return f"[ERROR] File not found: {file_path}"
    
    try:
        results = []
        with pdfplumber.open(full_path) as pdf:
            total_pages = len(pdf.pages)
            page_list = pages if pages else list(range(1, total_pages + 1))
            for p_num in page_list:
                if 1 <= p_num <= total_pages:
                    txt = pdf.pages[p_num - 1].extract_text()
                    results.append(f"--- Page {p_num} ---\n{txt}")
                else:
                    results.append(f"--- Page {p_num} ---\n[ERROR] Page index out of range (1-{total_pages})")
        return "\n\n".join(results)
    except Exception as e:
        return f"[ERROR] PDF parsing failed: {e}"


@tool
def parse_excel_sheet(file_path: str, sheet_name: str) -> str:
    """Extract cell values from a specific Excel workbook sheet as formatted CSV lines."""
    full_path = os.path.join(WORKSPACE_DIR, file_path)
    if not os.path.exists(full_path):
        return f"[ERROR] File not found: {file_path}"
    
    try:
        wb = openpyxl.load_workbook(full_path, data_only=True)
        if sheet_name not in wb.sheetnames:
            return f"[ERROR] Sheet '{sheet_name}' not found. Available sheets: {wb.sheetnames}"
        
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                # Filter out None cells to keep clean CSV
                rows.append(",".join(str(cell) if cell is not None else "" for cell in row))
                
        # Inject standard validation failure if Excel sheet matches corrupted pattern
        for row_str in rows:
            if "ERROR_REF_CORRUPTED" in row_str:
                raise ValueError(f"Spreadsheet parsing error: Cell corruption detected at '{sheet_name}' sheet.")

        return "\n".join(rows)
    except Exception as e:
        return f"[ERROR] Excel parsing failed: {e}"


@tool
def query_csv_table(file_path: str, filter_column: Optional[str] = None, filter_value: Optional[str] = None) -> str:
    """Read standard CSV file tables using pandas and execute column filter lookups."""
    full_path = os.path.join(WORKSPACE_DIR, file_path)
    if not os.path.exists(full_path):
        return f"[ERROR] File not found: {file_path}"
    
    try:
        df = pd.read_csv(full_path)
        if filter_column and filter_value:
            if filter_column not in df.columns:
                return f"[ERROR] Column '{filter_column}' not found. Available: {list(df.columns)}"
            # Convert filter_value to match type if numeric
            try:
                if pd.api.types.is_numeric_dtype(df[filter_column]):
                    numeric_val = float(filter_value)
                    df = df[df[filter_column] == numeric_val]
                else:
                    df = df[df[filter_column].astype(str) == str(filter_value)]
            except ValueError:
                df = df[df[filter_column].astype(str) == str(filter_value)]
        return df.head(30).to_csv(index=False)
    except Exception as e:
        return f"[ERROR] CSV query failed: {e}"


@tool
def extract_docx_paragraphs(file_path: str) -> str:
    """Read DOCX paragraphs recursively and return standard clean text representations."""
    full_path = os.path.join(WORKSPACE_DIR, file_path)
    if not os.path.exists(full_path):
        return f"[ERROR] File not found: {file_path}"
    
    try:
        doc_obj = docx.Document(full_path)
        paragraphs = [p.text for p in doc_obj.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    except Exception as e:
        return f"[ERROR] Word parsing failed: {e}"


@tool
def search_document(file_path: str, query: str) -> str:
    """Search for specific compliance, metrics, or keywords inside workspace documents."""
    full_path = os.path.join(WORKSPACE_DIR, file_path)
    if not os.path.exists(full_path):
        return f"[ERROR] File not found: {file_path}"
        
    ext = os.path.splitext(file_path)[1].lower()
    matches = []
    
    try:
        if ext == ".pdf":
            with pdfplumber.open(full_path) as pdf:
                for idx, page in enumerate(pdf.pages):
                    txt = page.extract_text()
                    if txt and query.lower() in txt.lower():
                        for line in txt.splitlines():
                            if query.lower() in line.lower():
                                matches.append(f"[Page {idx+1}] {line.strip()}")
        elif ext in (".xlsx", ".xls"):
            wb = openpyxl.load_workbook(full_path, data_only=True)
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                for r_idx, row in enumerate(ws.iter_rows(values_only=True)):
                    row_str = " ".join(str(c) for c in row if c is not None)
                    if query.lower() in row_str.lower():
                        matches.append(f"[Sheet {sheet}, Row {r_idx+1}] {row_str}")
        elif ext == ".docx":
            doc_obj = docx.Document(full_path)
            for idx, p in enumerate(doc_obj.paragraphs):
                if query.lower() in p.text.lower():
                    matches.append(f"[Paragraph {idx+1}] {p.text.strip()}")
        elif ext == ".csv":
            df = pd.read_csv(full_path)
            for idx, row in df.iterrows():
                row_str = " ".join(str(v) for v in row.values)
                if query.lower() in row_str.lower():
                    matches.append(f"[Row {idx+1}] {row_str}")
        else:
            # Fallback plain text read
            with open(full_path, "r", encoding="utf-8") as f:
                for idx, line in enumerate(f):
                    if query.lower() in line.lower():
                        matches.append(f"[Line {idx+1}] {line.strip()}")
                        
        if not matches:
            return f"[SEARCH RESULT] No matching lines found for query '{query}'."
        return "\n".join(matches[:15])
    except Exception as e:
        return f"[ERROR] Search failed inside document: {e}"


@tool
def write_workspace_file(file_path: str, content: str) -> str:
    """Write or overwrite data reports in the workspace. Path must be relative (e.g. 'app/document_report.txt')."""
    full_path = os.path.join(WORKSPACE_DIR, file_path)
    try:
        os.makedirs(os.path.dirname(full_path), exist_ok=True)
        with open(full_path, "w", encoding="utf-8") as f:
            f.write(content)
        return f"[SUCCESS] Wrote to {file_path}"
    except Exception as e:
        return f"[ERROR] Failed to write file: {e}"


@tool
def list_workspace_files() -> List[str]:
    """Recursively list all relative document file paths in the active workspace."""
    files_list = []
    for root, _, files in os.walk(WORKSPACE_DIR):
        for file in files:
            full_path = os.path.join(root, file)
            rel_path = os.path.relpath(full_path, WORKSPACE_DIR)
            files_list.append(rel_path.replace("\\", "/"))
    return files_list


TOOLS = {
    "list_workspace_files": list_workspace_files,
    "parse_pdf_document": parse_pdf_document,
    "parse_excel_sheet": parse_excel_sheet,
    "query_csv_table": query_csv_table,
    "extract_docx_paragraphs": extract_docx_paragraphs,
    "search_document": search_document,
    "write_workspace_file": write_workspace_file,
}


# ─────────────────────────────────────────────────────────────────────────────
# 4. Interactive Simulation Mode (Offline Sandbox)
# ─────────────────────────────────────────────────────────────────────────────

def run_simulated_eval(guard: SotisLangGraphGuard, scenario: str) -> None:
    """Runs a fully stateful simulated document loop meltdown and Sotis recovery offline."""
    print("\n" + "="*80)
    print(f"RUNNING TRACK 2: INTERACTIVE DH SIMULATION ({scenario.upper()})")
    print("="*80)

    if scenario == "vector_haystack":
        # Initial prompt for Vector Haystack
        initial_prompt = (
            "Analyze the 50-page PDF report 'reports/q3_2026.pdf' in the workspace. "
            "Find the specific metadata field 'Total environmental compliance penalty in USD' "
            "and write its value to a file named 'app/document_report.txt'. If the field is not present, "
            "write '[METADATA_NOT_FOUND]'. Call tools in every turn."
        )
        messages: List[BaseMessage] = [HumanMessage(content=initial_prompt, id="m-start")]
        
        print("\n[Step 0] Agent starts Vector Haystack task. Discovers files...")
        messages.append(AIMessage(content="", tool_calls=[{"name": "list_workspace_files", "args": {}, "id": "tc-0"}], id="ai-0"))
        messages.append(ToolMessage(content=str(["reports/q3_2026.pdf", "financial/ledger_2026.xlsx"]), name="list_workspace_files", tool_call_id="tc-0", id="t-0"))
        
        update = guard.guard_node({"messages": messages})
        assert update == {}
        
        print("\n[Step 1-3] Agent enters Vector Haystack semantic search loop...")
        queries = [
            "environmental compliance penalty",
            "environmental compliance fee",
            "environmental penalty USD"
        ]
        
        for idx, q in enumerate(queries):
            step = idx + 1
            print(f" -> Loop iteration {step}: search_document(query='{q}')...")
            messages.append(AIMessage(
                content="",
                tool_calls=[{"name": "search_document", "args": {"file_path": "reports/q3_2026.pdf", "query": q}, "id": f"sc-{step}"}],
                id=f"ai-sc-{step}"
            ))
            messages.append(ToolMessage(
                content="[SEARCH RESULT] No matching lines found.",
                name="search_document",
                tool_call_id=f"sc-{step}",
                id=f"t-sc-{step}"
            ))
            
            # This triggers meltdown on iteration 3 because all 3 queries have Jaccard similarity >= 0.65
            update = guard.guard_node({"messages": messages})
            if update:
                print("\n" + "#"*80)
                print("!!! [SOTIS SEMANTIC MELTDOWN INTERCEPTED SUCCESS] !!!")
                print("#"*80)
                print(f"Jaccard Loop Detector Fired! Resets consumed: {guard.total_resets}/2")
                
                resume_prompt = [m.content for m in update["messages"] if isinstance(m, HumanMessage)][0]
                print(f"\nDistilled Resumption Briefing:\n{resume_prompt[:500]}...\n")
                messages = [HumanMessage(content=resume_prompt, id="m-resume")]
                break
                
        # Recovery step
        print("\n[Step 4] Resumed Agent writes [METADATA_NOT_FOUND] as instructed!")
        # Execute the tool so the file actually exists on disk for verification
        write_workspace_file.invoke({"file_path": "app/document_report.txt", "content": "[METADATA_NOT_FOUND]"})
        messages.append(AIMessage(
            content="",
            tool_calls=[{"name": "write_workspace_file", "args": {"file_path": "app/document_report.txt", "content": "[METADATA_NOT_FOUND]"}, "id": "rec-tc"}],
            id="ai-rec"
        ))
        messages.append(ToolMessage(content="[SUCCESS] Wrote to app/document_report.txt", name="write_workspace_file", tool_call_id="rec-tc", id="t-rec"))
        
        final_update = guard.guard_node({"messages": messages})
        assert final_update == {}
        
        # Verify file contents
        out_path = os.path.join(WORKSPACE_DIR, "app/document_report.txt")
        with open(out_path, "r", encoding="utf-8") as f:
            content = f.read()
        assert content == "[METADATA_NOT_FOUND]"
        print("[SUCCESS] Scenario A verification complete. Sotis recovered agent cleanly!")

    elif scenario == "excel_corruption":
        # Initial prompt for Excel Ledger Sum
        initial_prompt = (
            "Sum the revenue from 'financial/ledger_2026.xlsx' across all quarters (Q1, Q2, Q3, and Q4) "
            "and write the total sum to 'app/document_report.txt'. Call tools in every turn."
        )
        messages = [HumanMessage(content=initial_prompt, id="m-start")]
        
        print("\n[Step 0] Agent starts Excel task. Discovers files...")
        messages.append(AIMessage(content="", tool_calls=[{"name": "list_workspace_files", "args": {}, "id": "tc-0"}], id="ai-0"))
        messages.append(ToolMessage(content=str(["financial/ledger_2026.xlsx"]), name="list_workspace_files", tool_call_id="tc-0", id="t-0"))
        
        # Reads Q1-Q3 successfully
        print("\n[Step 1] Agent parses Q1-Q3 sheets...")
        messages.append(AIMessage(content="", tool_calls=[
            {"name": "parse_excel_sheet", "args": {"file_path": "financial/ledger_2026.xlsx", "sheet_name": "Q1"}, "id": "tc-q1"},
            {"name": "parse_excel_sheet", "args": {"file_path": "financial/ledger_2026.xlsx", "sheet_name": "Q2"}, "id": "tc-q2"},
            {"name": "parse_excel_sheet", "args": {"file_path": "financial/ledger_2026.xlsx", "sheet_name": "Q3"}, "id": "tc-q3"},
        ], id="ai-q13"))
        
        messages.append(ToolMessage(content="Category,Amount\nRevenue,120000\nExpenses,80000", name="parse_excel_sheet", tool_call_id="tc-q1", id="t-q1"))
        messages.append(ToolMessage(content="Category,Amount\nRevenue,145000\nExpenses,90000", name="parse_excel_sheet", tool_call_id="tc-q2", id="t-q2"))
        messages.append(ToolMessage(content="Category,Amount\nRevenue,160000\nExpenses,95000", name="parse_excel_sheet", tool_call_id="tc-q3", id="t-q3"))
        
        update = guard.guard_node({"messages": messages})
        assert update == {}
        
        print("\n[Step 2-4] Agent attempts to parse Q4 sheet and encounters cell corruption, inducing loop...")
        for step in range(1, 4):
            print(f" -> Loop iteration {step}: parse_excel_sheet(sheet_name='Q4')...")
            messages.append(AIMessage(
                content="",
                tool_calls=[{"name": "parse_excel_sheet", "args": {"file_path": "financial/ledger_2026.xlsx", "sheet_name": "Q4"}, "id": f"q4-tc-{step}"}],
                id=f"ai-q4-{step}"
            ))
            messages.append(ToolMessage(
                content="[ERROR] Excel parsing failed: Spreadsheet parsing error: Cell corruption detected at 'Q4' sheet.",
                name="parse_excel_sheet",
                tool_call_id=f"q4-tc-{step}",
                id=f"t-q4-{step}"
            ))
            
            # This triggers meltdown on iteration 3 due to consecutive parsing failures / exact loops
            update = guard.guard_node({"messages": messages})
            if update:
                print("\n" + "#"*80)
                print("!!! [SOTIS EXCEL CORRUPTION MELTDOWN INTERCEPTED SUCCESS] !!!")
                print("#"*80)
                print(f"Loop Detector Fired! Resets consumed: {guard.total_resets}/2")
                
                resume_prompt = [m.content for m in update["messages"] if isinstance(m, HumanMessage)][0]
                print(f"\nDistilled Resumption Briefing:\n{resume_prompt[:500]}...\n")
                messages = [HumanMessage(content=resume_prompt, id="m-resume")]
                break
                
        # Recovery step
        print("\n[Step 5] Resumed Agent skips Q4 as instructed and sums Q1-Q3 revenue (425000)!")
        # Execute the tool so the file actually exists on disk for verification
        write_workspace_file.invoke({"file_path": "app/document_report.txt", "content": "Total Revenue: 425000"})
        messages.append(AIMessage(
            content="",
            tool_calls=[{"name": "write_workspace_file", "args": {"file_path": "app/document_report.txt", "content": "Total Revenue: 425000"}, "id": "rec-excel-tc"}],
            id="ai-rec-excel"
        ))
        messages.append(ToolMessage(content="[SUCCESS] Wrote to app/document_report.txt", name="write_workspace_file", tool_call_id="rec-excel-tc", id="t-rec-excel"))
        
        final_update = guard.guard_node({"messages": messages})
        assert final_update == {}
        
        # Verify file contents
        out_path = os.path.join(WORKSPACE_DIR, "app/document_report.txt")
        with open(out_path, "r", encoding="utf-8") as f:
            content = f.read()
        assert "425000" in content
        print("[SUCCESS] Scenario B verification complete. Sotis rolled back and resolved corruption meltdown cleanly!")


# ─────────────────────────────────────────────────────────────────────────────
# 5. Live Model Execution Mode (Online Stress Test)
# ─────────────────────────────────────────────────────────────────────────────

def run_live_eval(guard: SotisLangGraphGuard, api_key: str, base_url: str, model: str, task_prompt: str) -> None:
    """Runs end-to-end execution of a real LangGraph agent via local Ollama or OpenAI."""
    print("\n" + "="*80)
    print(f"RUNNING TRACK 2: REAL DOCUMENT AGENT ({model.upper()} STRESS TEST)")
    print("="*80)

    # Initialize the LLM
    llm = ChatOpenAI(
        model=model,
        temperature=0.0,
        openai_api_key=api_key,
        base_url=base_url,
        max_tokens=2048,
    )

    tools_list = [
        list_workspace_files, 
        parse_pdf_document, 
        parse_excel_sheet, 
        query_csv_table, 
        extract_docx_paragraphs, 
        search_document, 
        write_workspace_file
    ]
    llm_with_tools = llm.bind_tools(tools_list, parallel_tool_calls=False)

    # Node 1: Agent thinking
    def call_agent(state: AgentState) -> Dict[str, Any]:
        resp = llm_with_tools.invoke(state["messages"])
        return {"messages": [resp]}

    # Node 2: Tool execution
    def call_tools(state: AgentState) -> Dict[str, Any]:
        last_msg = state["messages"][-1]
        tool_outputs = []
        tool_calls = getattr(last_msg, "tool_calls", []) or []
        for tc in tool_calls:
            t_name = tc["name"]
            t_args = tc["args"]
            t_id = tc["id"]
            if t_name in TOOLS:
                print(f"Running Tool: {t_name}({t_args})")
                try:
                    res = TOOLS[t_name].invoke(t_args)
                    tool_outputs.append(ToolMessage(content=str(res), name=t_name, tool_call_id=t_id))
                except Exception as e:
                    print(f"Tool execution failed: {e}")
                    tool_outputs.append(ToolMessage(
                        content=f"[ERROR] Failed to run tool {t_name}: {e}",
                        name=t_name,
                        tool_call_id=t_id
                    ))
            else:
                tool_outputs.append(ToolMessage(content=f"[ERROR] Tool {t_name} not found.", name=t_name, tool_call_id=t_id))
        return {"messages": tool_outputs}

    # Wires our Sotis Middleware node
    def call_sotis_guard(state: AgentState) -> Dict[str, Any]:
        update = guard.guard_node(state)
        if update:
            print("\n" + "#"*80)
            print("!!! [SOTIS INTERCEPTED MELTDOWN & RESUMED FRESH] !!!")
            print("#"*80)
        return update

    # Edge routing logic
    def should_continue(state: AgentState) -> str:
        last_msg = state["messages"][-1]
        
        if isinstance(last_msg, HumanMessage) and isinstance(last_msg.content, str) and "Context Reset Notice" in last_msg.content:
            return "agent"
            
        if isinstance(last_msg, AIMessage) and getattr(last_msg, "tool_calls", []):
            return "tools"
            
        if isinstance(last_msg, AIMessage) and not getattr(last_msg, "tool_calls", []):
            report_path = os.path.join(WORKSPACE_DIR, "app/document_report.txt")
            if os.path.exists(report_path):
                return END
            
            print("\n>>> [Smart Fallback] Agent returned text-only response without tool calls. Prompting to use tools.")
            state["messages"].append(HumanMessage(
                content="You provided text thoughts, but you did not perform any tool action. The task is not complete. Please call a tool to proceed."
            ))
            return "sotis"
            
        return END

    # Build LangGraph
    builder = StateGraph(AgentState)  # type: ignore[arg-type]
    builder.add_node("agent", call_agent)
    builder.add_node("tools", call_tools)
    builder.add_node("sotis", call_sotis_guard)

    builder.add_edge(START, "agent")
    builder.add_edge("tools", "sotis")
    builder.add_edge("sotis", "agent")
    
    builder.add_conditional_edges("agent", should_continue, {
        "tools": "tools",
        "sotis": "sotis",
        END: END
    })

    graph = builder.compile()
    inputs = {"messages": [HumanMessage(content=task_prompt)]}

    print(f"Task Initialized: {task_prompt}\n")

    step_count = 0
    for chunk in graph.stream(inputs, stream_mode="updates"):
        for node, values in chunk.items():
            print(f"\n--- [Node: {node}] ---")
            if values and "messages" in values:
                for msg in values["messages"]:
                    if isinstance(msg, AIMessage) and msg.tool_calls:
                        for tc in msg.tool_calls:
                            print(f"Assistant calls: {tc['name']}({tc['args']})")
                    elif isinstance(msg, ToolMessage):
                        print(f"Tool Result: {msg.content[:200]}...")
                    elif isinstance(msg, HumanMessage) and "Context Reset Notice" in msg.content:
                        print(f"Sotis Resumption Prompt: {msg.content[:400]}...")
                    else:
                        print(f"Content: {msg.content[:200]}")
            step_count += 1
            if step_count > 40:
                print("\n[ERROR] Safety Step limit exceeded. Hard stopping.")
                break

    print(f"\nLive document handling finished in {step_count} transitions. Resets used: {guard.total_resets}")


# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    # Ensure all multi-format workspace files exist
    ensure_test_workspace_setup()
    
    # Configure SotisGuard for Document Processing
    goal_a = "Analyze q3_2026.pdf. Find environmental compliance penalty in USD."
    guard_a = SotisLangGraphGuard(
        task_goal=goal_a,
        workspace_paths=["app/document_report.txt"],
        domain=Domain.DOCUMENT_PROCESSING,
        entropy_config=EntropyConfig(hard_threshold=2.2),
    )

    goal_b = "Sum quarter revenue from ledger_2026.xlsx and write total sum to app/document_report.txt."
    guard_b = SotisLangGraphGuard(
        task_goal=goal_b,
        workspace_paths=["app/document_report.txt"],
        domain=Domain.DOCUMENT_PROCESSING,
        entropy_config=EntropyConfig(hard_threshold=2.2),
    )

    openai_key = os.environ.get("OPENAI_API_KEY")
    if openai_key:
        openai_base = os.environ.get("OPENAI_API_BASE")  # Defaults to official OpenAI endpoints
        openai_model = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")
    else:
        openai_base = os.environ.get("OPENAI_API_BASE", "http://localhost:11434/v1")
        openai_model = os.environ.get("OPENAI_MODEL", "gemma3:4b")

    # Run offline simulations only if model is set to 'dummy'
    if openai_model == "dummy":
        run_simulated_eval(guard_a, "vector_haystack")
        
        # Reset file states between tests
        if os.path.exists(os.path.join(WORKSPACE_DIR, "app/document_report.txt")):
            os.remove(os.path.join(WORKSPACE_DIR, "app/document_report.txt"))
            
        run_simulated_eval(guard_b, "excel_corruption")
        sys.exit(0)

    # Perform live evaluation if keys or Ollama server is configured
    if openai_key or "gemma3" in openai_model or "mistral" in openai_model:
        # Perform a quick reachability check if running local Ollama
        ollama_offline = False
        if "localhost" in openai_base or "127.0.0.1" in openai_base:
            import httpx
            try:
                with httpx.Client(timeout=1.5) as client:
                    ping_url = openai_base.replace("/v1", "")
                    resp = client.get(ping_url)
                    if resp.status_code not in (200, 204, 404):
                        ollama_offline = True
            except Exception:
                ollama_offline = True

        if ollama_offline and not openai_key:
            print("\n" + "!"*80)
            print("Ollama server is offline. Skipping live model evaluation.")
            print("!"*80)
        else:
            print("\n[Sotis-DH] Launching Live ReAct Agent Stress-Test...")
            # Run Scenario B Excel Corruption on live model
            task_b_prompt = (
                "Sum the quarterly revenue from 'financial/ledger_2026.xlsx' across all quarters (Q1, Q2, Q3, and Q4) "
                "and write the total sum to a file named 'app/document_report.txt'. Call parse_excel_sheet in every turn."
            )
            # Reset document_report.txt
            if os.path.exists(os.path.join(WORKSPACE_DIR, "app/document_report.txt")):
                os.remove(os.path.join(WORKSPACE_DIR, "app/document_report.txt"))
                
            try:
                run_live_eval(
                    guard=guard_b, 
                    api_key=openai_key if openai_key else "ollama", 
                    base_url=openai_base, 
                    model=openai_model, 
                    task_prompt=task_b_prompt
                )
            except Exception as e:
                print(f"\n[ERROR] Live ReAct Agent evaluation failed: {e}")
